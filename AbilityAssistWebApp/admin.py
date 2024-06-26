from django.contrib import admin
from django.contrib.auth.models import User
from django.contrib.auth.admin import UserAdmin as BaseUserAdmin
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.pdfmetrics import stringWidth

from .models import Trip, FinalGeolocation, InitialGeolocation, AboutImage, LocationUpdate
import csv
from reportlab.pdfgen import canvas
from docx import Document
from io import BytesIO
from django.http import HttpResponse
from django.utils.translation import gettext_lazy as _
from django.db.models import Q


class DistanceFilter(admin.SimpleListFilter):
    title = _('distance')
    parameter_name = 'distance'

    def lookups(self, request, model_admin):
        # This provides the options you see in the admin.
        options = [('0km', _('0 km'))]
        for i in range(1, 31):
            options.append((f'{i / 10}km', _(f'{i / 10} km')))
        options.append(('3.0km+', _('out of range trips')))
        return options

    def queryset(self, request, queryset):
        # This filters the queryset based on the selected option.
        if self.value() == '3.0km+':
            return queryset.filter(distance__gt=3.0).order_by('-distance')
        elif self.value():
            value_km = float(self.value().replace('km', ''))
            if value_km < 0.1:
                # Convert kilometers to meters and filter
                value_m = int(value_km * 1000)
                return queryset.filter(distance__lte=f'{value_m} m').order_by('-distance')
            else:
                return queryset.filter(distance__lte=value_km).order_by('-distance')
        else:
            return queryset.order_by('-distance')

class DurationFilter(admin.SimpleListFilter):
    title = _('duration')
    parameter_name = 'duration'

    def lookups(self, request, model_admin):
        # This provides the options you see in the admin.
        options = [('1min', _('1 min'))]
        for i in range(2, 21):
            options.append((f'{i}min', _(f'{i} mins')))
        options.append(('20min+', _('out of bound trips')))
        return options

    def queryset(self, request, queryset):
        # This filters the queryset based on the selected option.
        if self.value() == '20min+':
            query = Q()
            for i in range(21, 1000):  # Assuming a reasonable upper limit for duration
                query |= Q(duration=f"{i} mins")  # Look for exact string match
            return queryset.filter(query).order_by('-duration')
        elif self.value():
            duration_in_min = self.value().replace('min', '').strip()
            query = Q(duration__exact="1 min")  # Include the 1-minute trips
            for i in range(2, int(duration_in_min) + 1):
                query |= Q(duration=f"{i} mins")  # Look for exact string match
            return queryset.filter(query).order_by('-duration')
        return queryset.order_by('-duration')

# Define the export to CSV action
def export_to_csv(modeladmin, request, queryset):
    meta = modeladmin.model._meta
    field_names = [field.name for field in meta.fields]

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename={meta}.csv'
    writer = csv.writer(response)

    writer.writerow(field_names)
    for obj in queryset:
        row = writer.writerow([getattr(obj, field) for field in field_names])
    return response

export_to_csv.short_description = "Export Selected to CSV"


def export_to_pdf(modeladmin, request, queryset):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="selected_data.pdf"'
    p = canvas.Canvas(response, pagesize=letter)

    p.setFont("Helvetica-Bold", 14)
    model_name = modeladmin.model._meta.verbose_name_plural.title()
    p.drawString(100, 780, f"{model_name} Report")

    # Dynamic headers
    field_names = [field.verbose_name.title() for field in modeladmin.model._meta.fields]
    y_pos = 780
    p.setFont("Helvetica", 12)
    headers = " | ".join(field_names)
    p.drawString(100, y_pos, headers)

    def draw_wrapped_text(text, x, y, max_width):
        words = text.split()
        wrapped_lines = []
        current_line = ""
        for word in words:
            # Check if adding the next word exceeds the max width
            test_line = f"{current_line} {word}".strip()
            if stringWidth(test_line, "Helvetica", 10) <= max_width:
                current_line = test_line
            else:
                wrapped_lines.append(current_line)
                current_line = word
        wrapped_lines.append(current_line)  # Add the last line

        # Draw the lines
        for line in wrapped_lines:
            p.drawString(x, y, line)
            y -= 14  # Adjust line spacing
        return y  # Return the Y position after the last line

    p.setFont("Helvetica", 10)
    y_pos -= 30
    max_width = 500  # Maximum width for text before wrapping

    for obj in queryset:
        data_str = " | ".join([str(getattr(obj, field.name)) for field in modeladmin.model._meta.fields])
        y_pos = draw_wrapped_text(data_str, 100, y_pos, max_width)
        if y_pos < 100:
            p.showPage()
            y_pos = 800

    p.showPage()
    p.save()
    return response


export_to_pdf.short_description = "Export Selected to PDF"

def export_to_docx(modeladmin, request, queryset):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="selected_data.docx"'
    doc = Document()

    model_name = modeladmin.model._meta.verbose_name_plural.title()
    doc.add_heading(f'{model_name} Report', 0)

    table = doc.add_table(rows=1, cols=len(modeladmin.model._meta.fields))
    hdr_cells = table.rows[0].cells

    # Dynamically adjust column widths
    table.autofit = False
    for i, field in enumerate(modeladmin.model._meta.fields):
        hdr_cells[i].text = field.verbose_name.title()
        hdr_cells[i].width = Inches(1)  # Adjust the width as necessary

    # Table data
    for obj in queryset:
        row_cells = table.add_row().cells
        for i, field in enumerate(modeladmin.model._meta.fields):
            row_cells[i].text = str(getattr(obj, field.name))
            row_cells[i].width = Inches(1)  # Adjust the width as necessary

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    response.write(docx_stream.read())
    return response
export_to_docx.short_description = "Export Selected to DOCX"

def export_to_txt(modeladmin, request, queryset):
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename="selected_data.txt"'
    model_name = modeladmin.model._meta.verbose_name_plural.title()
    lines = [f"{model_name} Report\n"]

    # Dynamic headers
    field_names = [field.verbose_name.title() for field in modeladmin.model._meta.fields]
    lines.append(" | ".join(field_names) + "\n")

    # Data
    for obj in queryset:
        row = [str(getattr(obj, field.name)) for field in modeladmin.model._meta.fields]
        lines.append(" | ".join(row) + "\n")

    response.writelines(lines)
    return response
export_to_txt.short_description = "Export Selected to TXT"

# Extend the BaseUserAdmin class to include the export action
class UserAdmin(BaseUserAdmin):
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]
    search_fields = ['username', 'first_name', 'last_name', 'email']

    def mobile_numbers(self, obj):
        return obj.username

    mobile_numbers.short_description = 'Mobile Numbers'

    list_display = ('mobile_numbers', 'first_name', 'last_name', 'email', 'is_staff')

# ModelAdmin classes for your models
class TripAdmin(admin.ModelAdmin):
    list_display = ('user', 'start_point', 'end_point', 'date', 'distance', 'duration', 'cancelled', 'completed') # Fields to display in list view
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]
    search_fields = ['user__username']
    list_filter = ['cancelled', 'completed', 'date', DistanceFilter, DurationFilter]

class FinalGeolocationAdmin(admin.ModelAdmin):
    list_display = ('value', 'name')  # Fields to display in list view
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]
    search_fields = ['value', 'name']

class InitialGeolocationAdmin(admin.ModelAdmin):
    list_display = ('longitude', 'latitude', 'updated_at')  # Fields to display in list view
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]
    search_fields = ['longitude', 'latitude']

class LocationUpdateAdmin(admin.ModelAdmin):  # Define a new ModelAdmin class for LocationUpdate
    list_display = ('trip', 'latitude', 'longitude', 'timestamp')
    search_fields = ['trip__user__username']
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]

class AboutImageAdmin(admin.ModelAdmin):  # Define a new ModelAdmin class for AboutImage
    list_display = ('title', 'content')
    search_fields = ['title', 'content']
    actions = [export_to_csv, export_to_pdf, export_to_docx, export_to_txt]


# Unregister the original UserAdmin and register the new UserAdmin with export action
admin.site.unregister(User)
admin.site.register(User, UserAdmin)

# Register your models here
admin.site.register(Trip, TripAdmin)
admin.site.register(FinalGeolocation, FinalGeolocationAdmin)
admin.site.register(InitialGeolocation, InitialGeolocationAdmin)
admin.site.register(AboutImage, AboutImageAdmin)
admin.site.register(LocationUpdate, LocationUpdateAdmin)
