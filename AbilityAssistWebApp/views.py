# Standard library imports
import io
import json
from collections import defaultdict
from datetime import datetime, time

# Django imports
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.paginator import Paginator
from django.db.models import F, Count
from django.db.models import Window
from django.db.models.functions import DenseRank
from django.http import JsonResponse, HttpResponse, FileResponse, HttpResponseRedirect
from django.shortcuts import render, redirect
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
# External library imports for PDF and DOCX generation
from docx import Document
from reportlab.lib.pagesizes import letter

# Local application imports
from .forms import RegistrationForm, LoginForm, EditUserForm
from .models import Trip, InitialGeolocation, FinalGeolocation, AboutImage, LocationUpdate


def index(request):
    return render(request, 'index.html')


def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            user.save()
            profile = authenticate(request, username=user.username, password=user.password)
            login(request, user)
            return redirect('index')
    else:
        form = RegistrationForm()

    return render(request, 'register.html', {'form': form})

def user_login(request):
    if request.user.is_authenticated:
        # Redirect to the index page (or any other target page)
        return HttpResponseRedirect(reverse('index'))

    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user is not None:
                if user.is_staff:
                    login(request,user)
                    return redirect(reverse('admin:index'))
                login(request, user)
                return redirect('index')
            else:
                # Invalid credentials, display error message
                messages.error(request, 'Invalid username or password. Please try again. (If you have deactivated your account, please contact <a href="mailto:Abilityassistcompany@gmail.com">Abilityassistcompany@gmail.com</a> if you wish to re-activate it)')
                return render(request, 'login.html', {'form': form})
        else:
            # Invalid form submission, display error message
            messages.error(request, 'Invalid form submission. Please check your input.')
            return render(request, 'login.html', {'form': form})
    else:
        form = LoginForm()
        return render(request, 'login.html', {'form': form})



@login_required
def user_logout(request):
    logout(request)
    return redirect('login')

@login_required
def edit_profile(request):
    if request.method == 'POST':
        form = EditUserForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save(request=request,user=request.user)
            return redirect('index')
    else:
        form = EditUserForm(instance=request.user)
    return render(request, 'edit_profile.html', {'form': form})

@login_required
def deactivate_user(request):
    if request.method == "POST":
        user = request.user
        user.is_active = False  # Mark the user as inactive
        user.save()
        user_logout(request)
        # Redirect the user to the login page
        return redirect('login')
    else:
        return JsonResponse({"error": "Invalid request method"}, status=400)


@csrf_exempt
@login_required
@require_POST  # Ensures that this view can only be accessed by POST requests
def store_trip(request):
    # We no longer need to explicitly check if the method is POST due to the @require_POST decorator
    data = json.loads(request.body)
    initial_point_latitude = data['initialPoint']['latitude']
    initial_point_longitude = data['initialPoint']['longitude']
    final_destination_value = data['finalPoint']['finalDestinationValue']
    final_destination_name = data['finalPoint']['finalDestinationName']
    distance = data.get('distance', '')  # Capture distance
    duration = data.get('duration', '')  # Capture duration

    start_point_instance = InitialGeolocation.objects.create(
        latitude=initial_point_latitude,
        longitude=initial_point_longitude
    )

    end_point_instance, created = FinalGeolocation.objects.get_or_create(
        value=final_destination_value,
        name=final_destination_name
    )

    user_instance = request.user

    trip = Trip.objects.create(
        user=user_instance,
        start_point=start_point_instance,
        end_point=end_point_instance,
        distance=distance,  # Get the distance if provided
        duration=duration  # Get the duration if provided
    )

    response = {
        'trip_id': trip.id,
        'start_point': {
            'latitude': start_point_instance.latitude,
            'longitude': start_point_instance.longitude
        },
        'end_point': {
            'value': end_point_instance.value,
            'name': end_point_instance.name
        }
    }

    return JsonResponse(response, status=201)

@login_required
@require_POST
def update_current_location(request, trip_id):
    try:
        data = json.loads(request.body)
        trip = Trip.objects.get(id=trip_id, user=request.user)
        trip.start_point.latitude = data['latitude']
        trip.start_point.longitude = data['longitude']
        trip.start_point.save()

        # Update the trip's distance and duration
        LocationUpdate.objects.create(
            trip=trip,
            latitude=data['latitude'],
            longitude=data['longitude']
        )
        return JsonResponse({'success': 'Current location updated successfully'}, status=200)
    except Trip.DoesNotExist:
        return JsonResponse({'error': 'Trip not found'}, status=404)

    except KeyError as e:
        return JsonResponse({'error': f'Missing data: {e}'}, status=400)


@login_required
def trips(request):
    print("[DEBUG] Starting trips view")
    search_type = request.GET.get('searchType', 'no_filter')
    search_query = request.GET.get('search', '').strip()
    now = timezone.now()
    today_start = timezone.make_aware(datetime.combine(now.date(), datetime.min.time()))
    travel_entries = Trip.objects.filter(user=request.user).order_by('-id')

    print(f"[DEBUG] search_type: {search_type}, search_query: {search_query}")

    if search_type != 'no_filter' and search_query:
        try:
            if search_type == 'date':
                print("[DEBUG] Filtering by date")
                search_date = datetime.strptime(search_query, '%Y-%m-%d').date()
                travel_entries = travel_entries.filter(date__date=search_date)
            elif search_type == 'time':
                    print("[DEBUG] Filtering by time today")
                    try:
                        search_time = datetime.strptime(search_query, '%I:%M %p').time()
                        print("[DEBUG] Parsed search time (12-hour format):", search_time)
                    except ValueError:
                        search_time = datetime.strptime(search_query, '%H:%M').time()
                        print("[DEBUG] Parsed search time (24-hour format):", search_time)
                    end_of_day = timezone.make_aware(datetime.combine(now.date(), time.max))
                    print(f"[DEBUG] End of day: {end_of_day}")
                    travel_entries = travel_entries.filter(date__range=(today_start, end_of_day))
                    print(f"[DEBUG] Trips after filtering by day range: {travel_entries.count()}")
                    travel_entries_hour = travel_entries.filter(date__hour=search_time.hour)
                    print(f"[DEBUG] Trips after filtering by hour: {travel_entries_hour.count()}")
                    travel_entries_final = travel_entries_hour.filter(
                        date__minute__gte=max(0, search_time.minute - 15),
                        date__minute__lte=min(59, search_time.minute + 15)
                    )
                    print(f"[DEBUG] Trips after filtering by minute range: {travel_entries_final.count()}")
                    travel_entries = travel_entries_final
            elif search_type == 'datetime':
                print("[DEBUG] Filtering by datetime")
                search_query = search_query.replace(' PM', 'PM').replace(' AM', 'AM')
                try:
                    search_datetime = datetime.strptime(search_query, '%Y-%m-%d %I:%M%p')
                except ValueError:
                    search_datetime = datetime.strptime(search_query, '%Y-%m-%d %H:%M')
                search_datetime = timezone.make_aware(search_datetime)
                travel_entries = travel_entries.filter(
                    date__year=search_datetime.year,
                    date__month=search_datetime.month,
                    date__day=search_datetime.day,
                    date__hour=search_datetime.hour,
                    date__minute=search_datetime.minute
                )
            elif search_type == 'id':
                print("[DEBUG] Filtering by id")
                travel_entries = travel_entries.filter(id=int(search_query))

            elif search_type == 'destination':
                print("[DEBUG] Filtering by destination")
                travel_entries = travel_entries.filter(end_point__name__icontains=search_query)

            elif search_type == 'cancelled':
                print("[DEBUG] Filtering by cancellation")
                travel_entries = travel_entries.filter(cancelled=search_query.lower() == 'yes')

        except ValueError as e:
            print(f"[ERROR] Error parsing search query: {e}")
            messages.error(request, f"Error parsing search query: {e}")
            travel_entries = Trip.objects.none()

    request.session['filtered_trip_ids'] = list(travel_entries.values_list('id', flat=True))

    paginator = Paginator(travel_entries, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    print(f"[DEBUG] Total trips found: {paginator.count}")

    context = {
        'travel_entries': page_obj,
        'search_query': search_query,
        'search_type': search_type,
    }

    return render(request, 'trips.html', context)

@login_required
def trips_pdf(request):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="filtered_trips.pdf"'

    p = canvas.Canvas(response)

    filtered_trip_ids = request.session.get('filtered_trip_ids', [])
    trips = Trip.objects.filter(id__in=filtered_trip_ids).order_by('-date')

    y_pos = 800
    for trip in trips:
        p.drawString(100, y_pos,
                     f"Trip ID: {trip.id}, Destination: {trip.end_point.name}, Date: {trip.date.strftime('%Y-%m-%d %H:%M:%S')}")
        y_pos -= 20
        if y_pos < 100:
            p.showPage()
            y_pos = 800

    p.showPage()
    p.save()
    return response


@login_required
def trips_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="filtered_trips.csv"'

    writer = csv.writer(response)
    writer.writerow(['Trip ID', 'Start Location', 'Destination', 'Date', 'Cancelled'])

    filtered_trip_ids = request.session.get('filtered_trip_ids', [])
    trips = Trip.objects.filter(id__in=filtered_trip_ids).order_by('-date')

    for trip in trips:
        iso_formatted_date = trip.date.isoformat()
        writer.writerow([trip.id, f"{trip.start_point.latitude}, {trip.start_point.longitude}", trip.end_point.name,
                         iso_formatted_date, "Yes" if trip.cancelled else "No"])

    return response


@login_required
def trips_txt(request):
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename="filtered_trips.txt"'

    filtered_trip_ids = request.session.get('filtered_trip_ids', [])
    trips = Trip.objects.filter(id__in=filtered_trip_ids).order_by('-date')

    lines = []
    for trip in trips:
        lines.append(
            f"Trip ID: {trip.id}, Start Location: {trip.start_point.latitude}, {trip.start_point.longitude}, Destination: {trip.end_point.name}, Date: {trip.date.strftime('%Y-%m-%d %H:%M:%S')}, Cancelled: {'Yes' if trip.cancelled else 'No'}\n")

    response.writelines(lines)
    return response


@login_required
def trips_doc(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="filtered_trips.docx"'

    document = Document()
    document.add_heading('Trips', 0)

    filtered_trip_ids = request.session.get('filtered_trip_ids', [])
    trips = Trip.objects.filter(id__in=filtered_trip_ids).order_by('-date')

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Trip ID'
    hdr_cells[1].text = 'Start Location'
    hdr_cells[2].text = 'Destination'
    hdr_cells[3].text = 'Date'
    hdr_cells[4].text = 'Cancelled'

    for trip in trips:
        row_cells = table.add_row().cells
        row_cells[0].text = str(trip.id)
        row_cells[1].text = f"{trip.start_point.latitude}, {trip.start_point.longitude}"
        row_cells[2].text = trip.end_point.name
        row_cells[3].text = trip.date.strftime('%Y-%m-%d %H:%M:%S')
        row_cells[4].text = "Yes" if trip.cancelled else "No"

    docx_stream = io.BytesIO()
    document.save(docx_stream)
    docx_stream.seek(0)
    response.write(docx_stream.read())

    return response
@login_required
@require_POST
def update_final_destination(request, trip_id):
    data = json.loads(request.body)
    final_destination_value = data['finalDestinationValue']
    final_destination_name = data['finalDestinationName']

    end_point_instance, created = FinalGeolocation.objects.get_or_create(
        value=final_destination_value,
        name=final_destination_name
    )

    trip = Trip.objects.get(id=trip_id, user=request.user)
    trip.end_point = end_point_instance
    trip.cancelled = False
    trip.save()

    return JsonResponse({'success': 'Final destination updated successfully'}, status=200)


@login_required
@require_POST
def cancel_trip(request, trip_id):
    trip = Trip.objects.get(id=trip_id, user=request.user)
    trip.cancelled = True
    trip.completed = True
    trip.save()
    # Optionally, you can log the cancellation to the Django server console
    print(f'Trip ID {trip_id} cancelled successfully.')
    return JsonResponse({'success': 'Trip cancelled successfully'}, status=200)


@login_required
def recent_trip(request):
    try:
        latest_trip = Trip.objects.filter(user=request.user).latest('date')
    except Trip.DoesNotExist:
        latest_trip = None

    if 'pdf' in request.GET and latest_trip is not None:
        # Generate PDF response
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="recent_trip_{latest_trip.id}.pdf"'
        p = canvas.Canvas(response)
        p.drawString(100, 750, f"User: {request.user.username}")
        p.drawString(100, 735, f"Trip ID: {latest_trip.id}")
        p.drawString(100, 720, f"Start Location: {latest_trip.start_point.latitude}, {latest_trip.start_point.longitude}")
        p.drawString(100, 705, f"Destination: {latest_trip.end_point.name}")
        p.drawString(100, 690, f"Date: {latest_trip.date.strftime('%Y-%m-%d %H:%M:%S')}")
        p.drawString(100, 675, f"Cancelled: {'Yes' if latest_trip.cancelled else 'No'}")
        p.showPage()
        p.save()
        return response
    else:
        # HTML response
        context = {'latest_trip': latest_trip}
        return render(request, 'recent_trip.html', context)


# ... (other imports remain unchanged)

@login_required
def user_stats(request):
    print("[DEBUG] Starting user_stats view")
    search_query = request.GET.get('search', '')
    search_type = request.GET.get('search_type', 'all')
    print(f"[DEBUG] Search query: {search_query}, Search type: {search_type}")

    # Get users and annotate with trip count
    users = User.objects.annotate(trip_count=Count('trip')).order_by('-trip_count')

    # Apply DenseRank to all users
    users = users.annotate(rank=Window(expression=DenseRank(), order_by=F('trip_count').desc()))

    # Filter based on search type and query
    if search_type == 'username':
        users = users.filter(username__icontains=search_query)
    elif search_type == 'position' and search_query.isdigit():
        position = int(search_query)
        users = users.filter(rank=position)
    print(f"[DEBUG] Users found: {users.count()}")

    # Prepare data for downloads
    user_data_for_downloads = list(users.values('id', 'username', 'trip_count', 'rank'))
    request.session['user_data_for_downloads'] = user_data_for_downloads
    print("[DEBUG] Stored user data for downloads in session")

    # Paginate users
    paginator = Paginator(users, 10)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    print(f"[DEBUG] Page number: {page_number}")

    context = {
        'user_stats': page_obj,
        'search_query': search_query,
        'search_type': search_type,
    }

    print("[DEBUG] Prepared context for template")
    return render(request, 'user_stats.html', context)

def about(request):
    aboutImage = AboutImage.objects.last()
    context = {'aboutImage': aboutImage}
    return render(request, 'about.html', context)

def contact(request):
    if request.method == 'POST':
        pass
    return render(request, 'contact.html')


def help(request):
    return render(request,'help.html')

from reportlab.pdfgen import canvas
from io import BytesIO

@login_required
def recent_trip_pdf(request):
    try:
        latest_trip = Trip.objects.filter(user=request.user).latest('date')
    except Trip.DoesNotExist:
        latest_trip = None

    # Create a file-like buffer to receive PDF data
    buffer = BytesIO()

    # Create the PDF object, using the buffer as its "file."
    p = canvas.Canvas(buffer)

    # Draw things on the PDF. Here's where the PDF generation happens.
    if latest_trip:
        p.drawString(100, 800, f"Trip ID: {latest_trip.id}")
        p.drawString(100, 785, f"Start Location: {latest_trip.start_point.latitude}, {latest_trip.start_point.longitude}")
        p.drawString(100, 770, f"Destination: {latest_trip.end_point.name}")
        p.drawString(100, 755, f"Date: {latest_trip.date.strftime('%Y-%m-%d %H:%M:%S')}")
        p.drawString(100, 740, f"Cancelled: {'Yes' if latest_trip.cancelled else 'No'}")

    # Close the PDF object cleanly, and we're done.
    p.showPage()
    p.save()

    # File response
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename='recent_trip.pdf')

import csv

@login_required
def recent_trip_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="recent_trip.csv"'

    writer = csv.writer(response)
    writer.writerow(['Trip ID', 'Start Location', 'Destination', 'Date', 'Cancelled'])

    try:
        latest_trip = Trip.objects.filter(user=request.user).latest('date')
        iso_formatted_date = latest_trip.date.isoformat()
        writer.writerow([latest_trip.id, f"{latest_trip.start_point.latitude}, {latest_trip.start_point.longitude}", latest_trip.end_point.name, iso_formatted_date, "Yes" if latest_trip.cancelled else "No"])
    except Trip.DoesNotExist:
        pass  # Handle the case where there is no recent trip

    return response

@login_required
def recent_trip_txt(request):
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename="recent_trip.txt"'

    try:
        latest_trip = Trip.objects.filter(user=request.user).latest('date')
        response.write(f"Trip ID: {latest_trip.id}\nStart Location: {latest_trip.start_point.latitude}, {latest_trip.start_point.longitude}\nDestination: {latest_trip.end_point.name}\nDate: {latest_trip.date.strftime('%Y-%m-%d %H:%M:%S')}\nCancelled: {'Yes' if latest_trip.cancelled else 'No'}")
    except Trip.DoesNotExist:
        response.write("No recent trips found.")

    return response

@login_required
def recent_trip_doc(request):
    try:
        latest_trip = Trip.objects.filter(user=request.user).latest('date')
    except Trip.DoesNotExist:
        latest_trip = None

    # Create a Word document
    doc = Document()
    doc.add_heading('Recent Trip', 0)

    if latest_trip:
        p = doc.add_paragraph()
        p.add_run(f"Trip ID: {latest_trip.id}\n").bold = True
        p.add_run(f"Start Location: {latest_trip.start_point.latitude}, {latest_trip.start_point.longitude}\n")
        p.add_run(f"Destination: {latest_trip.end_point.name}\n")
        p.add_run(f"Date: {latest_trip.date.strftime('%Y-%m-%d %H:%M:%S')}\n")
        p.add_run(f"Cancelled: {'Yes' if latest_trip.cancelled else 'No'}")

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="recent_trip.docx"'

    return response

@login_required
def user_stats_pdf(request):
    print("[DEBUG] Generating user stats PDF")
    filtered_user_data = request.session.get('user_data_for_downloads', [])
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)

    p.drawString(72, 750, "Position")
    p.drawString(200, 750, "Username")
    p.drawString(400, 750, "Number of Trips")

    height = 730
    for user_data in filtered_user_data:
        p.drawString(72, height, str(user_data['rank']))
        p.drawString(200, height, user_data['username'])
        p.drawString(400, height, str(user_data['trip_count']))
        height -= 20
        if height < 100:
            p.showPage()
            height = 730

    p.showPage()
    p.save()
    buffer.seek(0)
    print("[DEBUG] Finished generating user stats PDF")
    return FileResponse(buffer, as_attachment=True, filename='user_stats.pdf')

@login_required
def user_stats_csv(request):
    print("[DEBUG] Generating user stats CSV")
    filtered_user_data = request.session.get('user_data_for_downloads', [])
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="user_stats.csv"'
    writer = csv.writer(response)
    writer.writerow(['Position', 'Username', 'Number of Trips'])

    for user_data in filtered_user_data:
        writer.writerow([user_data['rank'], user_data['username'], user_data['trip_count']])

    print("[DEBUG] Finished generating user stats CSV")
    return response

@login_required
def user_stats_txt(request):
    print("[DEBUG] Generating user stats TXT")
    filtered_user_data = request.session.get('user_data_for_downloads', [])
    response = HttpResponse(content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename="user_stats.txt"'

    lines = ["Position, Username, Number of Trips\n"]
    for user_data in filtered_user_data:
        lines.append(f"{user_data['rank']}, {user_data['username']}, {user_data['trip_count']}\n")

    response.writelines(lines)
    print("[DEBUG] Finished generating user stats TXT")
    return response

@login_required
def user_stats_doc(request):
    print("[DEBUG] Generating user stats DOCX")
    filtered_user_data = request.session.get('user_data_for_downloads', [])
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading('User Statistics', level=1)

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Position'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Number of Trips'

    for user_data in filtered_user_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(user_data['rank'])
        row_cells[1].text = user_data['username']
        row_cells[2].text = str(user_data['trip_count'])

    doc.save(buffer)
    buffer.seek(0)
    print("[DEBUG] Finished generating user stats DOCX")
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="user_stats.docx"'
    return response

@login_required
@require_POST
def complete_trip(request, trip_id):
    print(f"Received complete_trip request for trip ID: {trip_id}")
    try:
        trip = Trip.objects.get(id=trip_id, user=request.user)
        print(f"Trip found: {trip}")
        trip.completed = True
        trip.save()
        print(f"Trip ID {trip_id} marked as completed.")
        return JsonResponse({'success': True, 'message': 'Trip marked as completed.'})
    except Trip.DoesNotExist:
        print(f"Trip ID {trip_id} not found.")
        return JsonResponse({'success': False, 'error': 'Trip not found.'}, status=404)








